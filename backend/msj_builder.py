"""
MSJ Builder - Motion for Summary Judgment prompt engineering and document generation.

Handles:
- System prompt construction with no-hallucination guardrails
- Token budgeting for large document sets
- Citation verification post-processing
- DOCX export with proper legal document formatting
"""

import io
import re
from typing import Optional

# Lazy imports — loaded when DOCX functions are first called
_docx_imports_loaded = False
WD_ALIGN_PARAGRAPH = None
Inches = None
Pt = None
Twips = None
WD_TABLE_ALIGNMENT = None
qn = None
DocxDocument = None


def _ensure_docx_imports():
    global _docx_imports_loaded, WD_ALIGN_PARAGRAPH, Inches, Pt, Twips, WD_TABLE_ALIGNMENT, qn, DocxDocument
    if _docx_imports_loaded:
        return
    from docx import Document as _Doc
    from docx.shared import Inches as _Inches, Pt as _Pt, Twips as _Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD
    from docx.enum.table import WD_TABLE_ALIGNMENT as _WT
    from docx.oxml.ns import qn as _qn
    DocxDocument = _Doc
    Inches = _Inches
    Pt = _Pt
    Twips = _Twips
    WD_ALIGN_PARAGRAPH = _WD
    WD_TABLE_ALIGNMENT = _WT
    qn = _qn
    _docx_imports_loaded = True


# Token budget constants (Sonnet 200K context)
MAX_DOCUMENT_TOKENS = 120_000
MAX_LIBRARY_TOKENS = 15_000
MAX_RULE_TOKENS = 10_000
CHARS_PER_TOKEN = 4  # rough estimate


def _truncate_text(text: str, max_chars: int) -> str:
    """Truncate text preserving beginning and end."""
    if not text or len(text) <= max_chars:
        return text or ""
    # Keep first 80% and last 20%
    head = int(max_chars * 0.8)
    tail = max_chars - head
    return text[:head] + "\n\n[...document truncated...]\n\n" + text[-tail:]


def _budget_documents(documents: list[tuple], max_total_chars: int) -> list[tuple]:
    """
    Allocate character budget across documents.
    documents: [(id, doc_type, title, extracted_text)]
    Returns same structure with text potentially truncated.
    """
    if not documents:
        return []

    total_chars = sum(len(d[3] or "") for d in documents)
    if total_chars <= max_total_chars:
        return documents

    # Prioritize pleadings (step 2) over evidence (step 4)
    per_doc_budget = max_total_chars // len(documents)
    result = []
    for doc_id, doc_type, title, text in documents:
        # Give pleadings 50% more budget
        budget = int(per_doc_budget * 1.5) if doc_type == "pleading" else per_doc_budget
        result.append((doc_id, doc_type, title, _truncate_text(text, budget)))

    return result


def build_msj_system_prompt(
    case_info: dict,
    material_facts: list,
    legal_arguments: list,
    documents: list[tuple],
    library_docs: list[tuple],
    rule_56_text: Optional[str] = None,
    generated_motion: Optional[str] = None,
    core_cases: Optional[list[tuple]] = None,
) -> str:
    """
    Build the system prompt for MSJ chat interactions.

    Args:
        case_info: {plaintiff, defendant, court, jurisdiction, case_number, representing_side, judge}
        material_facts: [{fact_number, text, source_doc_id, source_excerpt}]
        legal_arguments: [{issue, standard, argument_text, supporting_case_ids, supporting_rule_ids}]
        documents: [(id, doc_type, title, extracted_text)]
        library_docs: [(title, content)]
        rule_56_text: Full text of FRCP Rule 56
        generated_motion: Previously generated motion text (if any)
    """
    parts = []

    parts.append("""You are a law school educational assistant helping a student draft a Motion for Summary Judgment.

CRITICAL RULES — YOU MUST FOLLOW THESE EXACTLY:
1. You may ONLY cite cases, rules, and statutes that appear in the APPROVED SOURCES section below.
2. You may ONLY reference facts that appear in the UPLOADED DOCUMENTS section or the STATEMENT OF UNDISPUTED MATERIAL FACTS.
3. NEVER fabricate, invent, or hallucinate any citation, case name, rule number, or factual claim.
4. If you are unsure whether a source exists in the provided materials, say so explicitly rather than guessing. If you need a legal principle but cannot find supporting authority in the approved sources, write "[CITATION NEEDED — verify before filing]" so the student knows to research it.
5. Every factual assertion must include a source reference like (Doc #X) or (Doc #X at [page]).
6. Every legal citation must use the EXACT case name and reporter citation from the approved sources. Do not paraphrase or alter citations. For Ohio cases, use the Ohio State Reports citation. For federal cases, use the U.S. Reports or Federal Reporter citation.
7. When citing a legal rule or standard, always include: (a) the specific rule or statute number, (b) the case that established or applied the rule, and (c) the exact citation from the approved sources.
8. This is an EDUCATIONAL tool for law students. Your output is for learning purposes only and must not be filed with any court.
9. If the student asks you to cite something not in the approved sources, explain that you can only work with provided materials and suggest they upload additional documents or add the source to the approved library.
10. For statutes and rules of procedure, always cite the specific section number (e.g., "Fed. R. Civ. P. 56(a)" not just "Rule 56").""")

    # Case information
    if case_info and any(case_info.values()):
        parts.append("\n\n--- CASE INFORMATION ---")
        if case_info.get("plaintiff"):
            parts.append(f"Plaintiff: {case_info['plaintiff']}")
        if case_info.get("defendant"):
            parts.append(f"Defendant: {case_info['defendant']}")
        if case_info.get("court"):
            parts.append(f"Court: {case_info['court']}")
        if case_info.get("jurisdiction"):
            parts.append(f"Jurisdiction: {case_info['jurisdiction']}")
        if case_info.get("case_number"):
            parts.append(f"Case Number: {case_info['case_number']}")
        if case_info.get("representing_side"):
            parts.append(f"Representing: {case_info['representing_side']}")
        if case_info.get("judge"):
            parts.append(f"Judge: {case_info['judge']}")

    # Uploaded documents (with token budgeting)
    if documents:
        budgeted_docs = _budget_documents(documents, MAX_DOCUMENT_TOKENS * CHARS_PER_TOKEN)
        parts.append("\n\n--- UPLOADED DOCUMENTS ---")
        parts.append("CITATION FORMAT: When citing these documents, use the format: (Last Name, Deposition, at [paragraph]) for depositions, or (Title, at [page]) for exhibits. Do NOT use 'Doc #' references. Use the document TITLE or deponent's LAST NAME.")
        for doc_id, doc_type, title, text in budgeted_docs:
            parts.append(f"\nDocument #{doc_id} ({doc_type}): {title}")
            parts.append(text or "[No text extracted]")

    # Statement of undisputed material facts
    if material_facts:
        parts.append("\n\n--- STATEMENT OF UNDISPUTED MATERIAL FACTS ---")
        for fact in material_facts:
            num = fact.get("fact_number", "?")
            text = fact.get("text", "")
            source_ref = ""
            source_ids = fact.get("source_doc_ids") or ([fact["source_doc_id"]] if fact.get("source_doc_id") else [])
            if source_ids:
                refs = ", ".join(f"Doc #{sid}" for sid in source_ids)
                source_ref = f" ({refs})"
            if fact.get("source_excerpt"):
                source_ref += f" [{fact['source_excerpt']}]"
            parts.append(f"{num}. {text}{source_ref}")

    # Legal arguments
    if legal_arguments:
        parts.append("\n\n--- LEGAL ARGUMENTS ---")
        for i, arg in enumerate(legal_arguments, 1):
            parts.append(f"\nArgument {i}: {arg.get('issue', 'Untitled')}")
            if arg.get("standard"):
                parts.append(f"Standard: {arg['standard']}")
            if arg.get("argument_text"):
                parts.append(f"Argument: {arg['argument_text']}")
            if arg.get("supporting_case_ids"):
                parts.append(f"Supporting cases: {', '.join(str(c) for c in arg['supporting_case_ids'])}")
            if arg.get("supporting_rule_ids"):
                parts.append(f"Supporting rules: {', '.join(str(r) for r in arg['supporting_rule_ids'])}")

    # Approved sources - Rule 56
    parts.append("\n\n--- APPROVED SOURCES ---")

    # Core summary judgment cases (always included, from database)
    if core_cases:
        parts.append("\nCore Summary Judgment Cases (USE THESE EXACT CITATIONS):")
        for title, reporter_cite, summary in core_cases:
            parts.append(f"\n{title}, {reporter_cite}")
            if summary:
                parts.append(f"Summary: {_truncate_text(summary, 3000)}")

    if rule_56_text:
        parts.append(f"\nFRCP Rule 56 (Summary Judgment):\n{_truncate_text(rule_56_text, MAX_RULE_TOKENS * CHARS_PER_TOKEN)}")

    # Approved sources - Library documents
    if library_docs:
        parts.append("\n\nLibrary Resources:")
        budget_per_lib = (MAX_LIBRARY_TOKENS * CHARS_PER_TOKEN) // max(len(library_docs), 1)
        for title, content in library_docs:
            parts.append(f"\n{title}:\n{_truncate_text(content, budget_per_lib)}")

    # Previously generated motion (for refinement)
    if generated_motion:
        parts.append(f"\n\n--- PREVIOUSLY GENERATED MOTION DRAFT ---\n{_truncate_text(generated_motion, 30000)}")

    return "\n".join(parts)


def build_msj_generation_prompt(
    case_info: dict,
    material_facts: list,
    legal_arguments: list,
    documents: list[tuple],
    library_docs: list[tuple],
    rule_56_text: Optional[str] = None,
    core_cases: Optional[list[tuple]] = None,
) -> tuple[str, str]:
    """
    Build system + user prompts specifically for generating the full motion document.
    Returns (system_prompt, user_prompt).
    """
    system_prompt = build_msj_system_prompt(
        case_info=case_info,
        material_facts=material_facts,
        legal_arguments=legal_arguments,
        documents=documents,
        library_docs=library_docs,
        rule_56_text=rule_56_text,
        core_cases=core_cases,
    )

    plaintiff = case_info.get("plaintiff", "[Plaintiff]")
    defendant = case_info.get("defendant", "[Defendant]")
    court = case_info.get("court", "[Court]")
    case_number = case_info.get("case_number", "[Case No.]")
    representing = case_info.get("representing_side", "plaintiff")

    user_prompt = f"""Generate the MEMORANDUM IN SUPPORT of a Motion for Summary Judgment for {representing} in {plaintiff} v. {defendant}, Case No. {case_number}, in {court}.

DO NOT generate the caption, motion title, short-form motion paragraph, signature block, or certificate of service — those are built separately. Generate ONLY the memorandum content, starting with the Introduction.

Structure with these sections using markdown headings:

## Introduction
Brief overview: what this case is about, what relief is sought, and why summary judgment is appropriate. (1-2 paragraphs)

## Statement of Undisputed Material Facts
Numbered facts, each with a citation to the record. Reference uploaded documents as (Doc #X) or (Doc #X at [page/paragraph]).

## Motion Standard
State the Rule 56(a) standard and applicable case law on burdens of production.

## Argument
Organize with ### subheadings for each legal issue or element challenged. For EACH issue, follow this exact structure:

1. STATE WHAT YOU WANT: Begin with what the court should conclude. Example: "Summary judgment should be granted because [opposing party] cannot produce sufficient evidence that [element]."
2. STATE THE LAW: Cite the governing substantive law for that element with binding authority and pin cites.
3. APPLY TO UNDISPUTED FACTS: Show how the specific undisputed facts (citing record evidence) satisfy or fail to satisfy the element. Weave the burden language throughout — remind the court who bears the burden and that they have not met it.
4. CONCLUDE WITH BURDEN: End each section by restating that the opposing party bears the burden and has failed to produce sufficient evidence for a rational fact-finder. Example: "Because [party] bears the burden of proof on [element] and has not produced sufficient evidence that a rational fact-finder could find in their favor, summary judgment is warranted on this issue."

Repeat this structure for EACH element or issue addressed.

## Conclusion
State the SPECIFIC relief requested. Be precise about what you want the court to do — do not ask for general relief. Example: "For the foregoing reasons, {representing.capitalize()} respectfully requests that this Court grant summary judgment in {representing.capitalize()}'s favor on Defendant's affirmative defense of [specific defense] and enter judgment accordingly."

IMPORTANT FORMATTING RULES:
- Use ## for main section headings and ### for argument sub-headings
- Number each fact in the Statement of Facts (1., 2., 3., etc.)
- Use standard legal citation format with PIN CITES (specific page numbers). Example: Celotex Corp. v. Catrett, 477 U.S. 317, 323 (1986)
- For deposition citations use EXACTLY this format: (Last Name, Deposition, at [paragraph number]) — for example: (Hayes, Deposition, at 3) or (Downing, Deposition, at 5). Use the deponent's LAST NAME, not "Doc #".
- For exhibit citations use: (Exhibit Title) — for example: (ZAMR Corp. Work Rules) or (Position Evaluation, Jan. 15, 2024)
- Do NOT use "Doc #" references anywhere. Always use the document title or deponent's last name.
- Write in formal legal prose appropriate for a federal court filing
- Every fact must cite to a specific document from the record using the formats above
- Every legal proposition must cite to an approved source with a pin cite
- Do NOT invent or fabricate any citation
- This is for educational purposes only

Generate the memorandum content now."""

    return system_prompt, user_prompt


def generate_motion_docx(case_info: dict, motion_text: str) -> bytes:
    """
    Generate a DOCX file from the motion text with proper federal court formatting.
    Builds the caption, signature block, and certificate of service from structured
    case_info data. The AI-generated motion_text fills the substantive sections.
    Returns the DOCX as bytes.
    """
    _ensure_docx_imports()

    doc = DocxDocument()

    # --- Page setup ---
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        # Page numbers - bottom center
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add page number field
        run = fp.add_run()
        fldChar1 = run._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
        run._r.append(fldChar1)
        run2 = fp.add_run()
        instrText = run2._r.makeelement(qn('w:instrText'), {})
        instrText.text = ' PAGE '
        run2._r.append(instrText)
        run3 = fp.add_run()
        fldChar2 = run3._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
        run3._r.append(fldChar2)
        for r in [run, run2, run3]:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)

    # --- Default style: TNR 12pt, double-spaced ---
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    plaintiff = case_info.get("plaintiff", "[Plaintiff]")
    defendant = case_info.get("defendant", "[Defendant]")
    court = case_info.get("court", "[United States District Court]")
    case_number = case_info.get("case_number", "[Case No.]")
    representing = case_info.get("representing_side", "plaintiff")
    judge = case_info.get("judge", "")

    # --- Educational disclaimer ---
    _add_centered_para(doc, "EDUCATIONAL DOCUMENT — NOT FOR FILING", bold=True, size=11)
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = dp.add_run(
        "This motion was generated by an AI tool for educational purposes only. "
        "It has not been reviewed by a licensed attorney and should not be filed with any court."
    )
    dr.italic = True
    dr.font.size = Pt(10)
    dr.font.name = "Times New Roman"
    doc.add_paragraph()  # spacer

    # --- Court name ---
    _add_centered_para(doc, "IN THE " + court.upper(), bold=True, size=12)
    doc.add_paragraph()

    # --- Caption block (using a table for the ) column layout) ---
    caption_table = doc.add_table(rows=5, cols=3)
    caption_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Remove all borders from table
    for row in caption_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.makeelement(qn('w:tcBorders'), {})
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = tcBorders.makeelement(qn(f'w:{border_name}'), {
                    qn('w:val'): 'none', qn('w:sz'): '0', qn('w:space'): '0', qn('w:color'): 'auto'
                })
                tcBorders.append(border)
            tcPr.append(tcBorders)

    # Set column widths
    for row in caption_table.rows:
        row.cells[0].width = Inches(3.0)
        row.cells[1].width = Inches(0.5)
        row.cells[2].width = Inches(3.0)

    # Fill in caption
    _set_cell_text(caption_table.cell(0, 0), plaintiff + ",", bold=False)
    _set_cell_text(caption_table.cell(0, 1), ")", bold=False)
    _set_cell_text(caption_table.cell(0, 2), "", bold=False)

    _set_cell_text(caption_table.cell(1, 0), "          Plaintiff,", bold=False)
    _set_cell_text(caption_table.cell(1, 1), ")", bold=False)
    _set_cell_text(caption_table.cell(1, 2), f"Case No. {case_number}", bold=True)

    _set_cell_text(caption_table.cell(2, 0), "     v.", bold=False)
    _set_cell_text(caption_table.cell(2, 1), ")", bold=False)
    _set_cell_text(caption_table.cell(2, 2), "", bold=False)

    _set_cell_text(caption_table.cell(3, 0), defendant + ",", bold=False)
    _set_cell_text(caption_table.cell(3, 1), ")", bold=False)
    _set_cell_text(caption_table.cell(3, 2), "", bold=False)

    _set_cell_text(caption_table.cell(4, 0), "          Defendant.", bold=False)
    _set_cell_text(caption_table.cell(4, 1), ")", bold=False)
    _set_cell_text(caption_table.cell(4, 2), "", bold=False)

    doc.add_paragraph()

    # --- Motion title ---
    movant = plaintiff if representing == "plaintiff" else defendant
    _add_centered_para(doc, f"MOTION OF {representing.upper()}, {movant.upper()},", bold=True, size=12)
    _add_centered_para(doc, "FOR SUMMARY JUDGMENT", bold=True, size=12)
    doc.add_paragraph()

    # --- Short-form motion paragraph ---
    short_motion = (
        f"{movant}, respectfully moves this Court for an Order granting summary judgment "
        f"in {movant}'s favor pursuant to Fed. R. Civ. P. 56(a) on the grounds that there is "
        f"no genuine dispute as to any material fact and {movant} is entitled to judgment as a "
        f"matter of law. A Memorandum in Support of this Motion is attached."
    )
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5)
    _add_formatted_run(p, short_motion)
    doc.add_paragraph()

    # --- Signature block for motion ---
    _add_signature_block(doc, representing, movant)

    # --- Page break before memorandum ---
    doc.add_page_break()

    # --- Memorandum title ---
    _add_centered_para(doc, "MEMORANDUM IN SUPPORT OF MOTION OF", bold=True, size=12)
    _add_centered_para(doc, f"{movant.upper()} FOR SUMMARY JUDGMENT", bold=True, size=12)
    doc.add_paragraph()

    # --- Parse and insert the AI-generated substantive content ---
    # Strip any caption/title the AI may have generated (we built it structurally above)
    content_lines = motion_text.split("\n")
    # Skip lines that look like AI-generated caption/title until we hit substance
    start_idx = 0
    for i, line in enumerate(content_lines):
        stripped = line.strip().lower()
        if any(kw in stripped for kw in ["introduction", "preliminary statement", "statement of",
                                          "i.", "1.", "## i", "## introduction"]):
            start_idx = i
            break
        # Also start at first markdown heading that's not a caption
        if line.strip().startswith("#") and not any(kw in stripped for kw in ["caption", "motion of", "court"]):
            start_idx = i
            break
    else:
        start_idx = 0  # If no marker found, use everything

    _render_motion_content(doc, content_lines[start_idx:])

    # --- Conclusion signature block ---
    doc.add_paragraph()
    _add_signature_block(doc, representing, movant)

    # --- Certificate of Service ---
    doc.add_page_break()
    _add_centered_para(doc, "CERTIFICATE OF SERVICE", bold=True, size=12)
    doc.add_paragraph()

    opposing = defendant if representing == "plaintiff" else plaintiff
    cos_text = (
        f"I hereby certify that a copy of the foregoing Motion for Summary Judgment and "
        f"Memorandum in Support was served by [method of service] on counsel for "
        f"{opposing}, at [address], this ___ day of __________, ____."
    )
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5)
    _add_formatted_run(p, cos_text)
    doc.add_paragraph()
    doc.add_paragraph()

    # Signature line for COS
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = sig.add_run("_________________________________")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = name_p.add_run(f"Attorney for {representing.capitalize()}")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)

    # --- Bottom disclaimer ---
    doc.add_paragraph()
    doc.add_paragraph()
    bd = doc.add_paragraph()
    bd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = bd.add_run(
        "EDUCATIONAL DOCUMENT — NOT FOR FILING\n"
        "Generated by Sage's Study Group (lawstudygroup.com) for educational purposes only."
    )
    br.italic = True
    br.font.size = Pt(10)
    br.font.name = "Times New Roman"

    # Save
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _set_cell_text(cell, text: str, bold: bool = False):
    """Set text in a table cell with TNR 12pt formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold


def _add_centered_para(doc, text: str, bold: bool = False, size: int = 12):
    """Add a centered paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return p


def _add_signature_block(doc, representing: str, movant: str):
    """Add a right-aligned signature block."""
    lines = [
        "Respectfully submitted,",
        "",
        "_________________________________",
        "[Attorney Name]",
        f"Attorney for {representing.capitalize()}, {movant}",
        "[Address]",
        "[City, State ZIP]",
        "[Phone]",
        "[Bar Number]",
    ]
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def _add_formatted_run(paragraph, text: str):
    """Add a run of text with TNR 12pt to a paragraph, handling inline bold/italic."""
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def _render_motion_content(doc, lines: list[str]):
    """
    Render the AI-generated motion body into the DOCX with proper legal formatting.
    Handles markdown headings, numbered lists, and body paragraphs.
    """
    for line in lines:
        stripped = line.strip()

        if not stripped:
            continue

        # Section headings (centered, bold, caps)
        if stripped.startswith("# "):
            _add_centered_para(doc, stripped[2:].strip().upper(), bold=True, size=12)
            continue

        # Sub-section headings (centered, bold)
        if stripped.startswith("## "):
            _add_centered_para(doc, stripped[3:].strip(), bold=True, size=12)
            continue

        # Sub-sub-headings (####) — left-aligned, bold, italic, indented
        if stripped.startswith("#### "):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(stripped[5:].strip())
            run.bold = True
            run.italic = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            continue

        # Argument sub-headings (left-aligned, bold, indented)
        if stripped.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(stripped[4:].strip())
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            continue

        # Bold-only line (often a section label)
        if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            _add_centered_para(doc, stripped[2:-2], bold=True, size=12)
            continue

        # Numbered fact lines (e.g., "1. Fact text here")
        num_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if num_match:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            _add_formatted_run(p, f"{num_match.group(1)}. {num_match.group(2)}")
            continue

        # Regular body paragraph with first-line indent
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.5)
        _add_formatted_run(p, stripped)


async def verify_citations(text: str, db_pool) -> dict:
    """
    Post-process AI-generated text to verify citations against the database.
    Returns {verified: [...], unverified: [...], text: modified_text}
    """
    try:
        from eyecite import get_citations
        from eyecite.models import CaseCitation
    except ImportError:
        return {"verified": [], "unverified": [], "text": text}

    citations = get_citations(text)
    verified = []
    unverified = []

    async with db_pool.acquire() as conn:
        for cite in citations:
            if not isinstance(cite, CaseCitation):
                continue

            cite_str = str(cite)
            # Try to find in database
            row = await conn.fetchrow(
                "SELECT id, title FROM cases WHERE reporter_cite ILIKE $1 LIMIT 1",
                f"%{cite_str}%"
            )
            if row:
                verified.append({"citation": cite_str, "case_id": row["id"], "title": row["title"]})
            else:
                unverified.append({"citation": cite_str})

    return {
        "verified": verified,
        "unverified": unverified,
        "text": text,
    }
