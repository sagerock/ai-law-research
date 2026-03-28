"""
Shared utilities for legal document builders.

Provides:
- Lazy docx imports
- Token budgeting for large document sets
- DOCX formatting helpers (caption, signature block, jurat, etc.)
- Markdown-to-DOCX rendering
- Citation verification
- Anti-hallucination prompt rules
"""

import io
import re
from typing import Optional

# Lazy imports — loaded when DOCX functions are first called
_docx_imports_loaded = False
WD_ALIGN_PARAGRAPH = None
Inches = None
Pt = None
Twips = None
WD_TABLE_ALIGNMENT = None
qn = None
DocxDocument = None


def ensure_docx_imports():
    global _docx_imports_loaded, WD_ALIGN_PARAGRAPH, Inches, Pt, Twips, WD_TABLE_ALIGNMENT, qn, DocxDocument
    if _docx_imports_loaded:
        return
    from docx import Document as _Doc
    from docx.shared import Inches as _Inches, Pt as _Pt, Twips as _Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD
    from docx.enum.table import WD_TABLE_ALIGNMENT as _WT
    from docx.oxml.ns import qn as _qn
    DocxDocument = _Doc
    Inches = _Inches
    Pt = _Pt
    Twips = _Twips
    WD_ALIGN_PARAGRAPH = _WD
    WD_TABLE_ALIGNMENT = _WT
    qn = _qn
    _docx_imports_loaded = True


# Token budget constants (Sonnet 200K context)
MAX_DOCUMENT_TOKENS = 120_000
MAX_LIBRARY_TOKENS = 15_000
MAX_RULE_TOKENS = 10_000
CHARS_PER_TOKEN = 4  # rough estimate


# ── Token budgeting ──────────────────────────────────────────────────────────

def truncate_text(text: str, max_chars: int) -> str:
    """Truncate text preserving beginning and end."""
    if not text or len(text) <= max_chars:
        return text or ""
    head = int(max_chars * 0.8)
    tail = max_chars - head
    return text[:head] + "\n\n[...document truncated...]\n\n" + text[-tail:]


def budget_documents(documents: list[tuple], max_total_chars: int, priority_type: str = "pleading") -> list[tuple]:
    """
    Allocate character budget across documents.
    documents: [(id, doc_type, title, extracted_text)]
    priority_type: doc_type that gets 50% more budget allocation.
    Returns same structure with text potentially truncated.
    """
    if not documents:
        return []

    total_chars = sum(len(d[3] or "") for d in documents)
    if total_chars <= max_total_chars:
        return documents

    per_doc_budget = max_total_chars // len(documents)
    result = []
    for doc_id, doc_type, title, text in documents:
        budget = int(per_doc_budget * 1.5) if doc_type == priority_type else per_doc_budget
        result.append((doc_id, doc_type, title, truncate_text(text, budget)))

    return result


# ── Prompt building helpers ──────────────────────────────────────────────────

def build_anti_hallucination_rules(tool_name: str = "legal document") -> str:
    """Core anti-hallucination rules shared across all tool types."""
    return f"""CRITICAL RULES — YOU MUST FOLLOW THESE EXACTLY:
1. You may ONLY cite cases, rules, and statutes that appear in the APPROVED SOURCES section below.
2. You may ONLY reference facts that appear in the UPLOADED DOCUMENTS section or the user-provided facts.
3. NEVER fabricate, invent, or hallucinate any citation, case name, rule number, or factual claim.
4. If you are unsure whether a source exists in the provided materials, say so explicitly rather than guessing. If you need a legal principle but cannot find supporting authority in the approved sources, write "[CITATION NEEDED — verify before filing]" so the student knows to research it.
5. Every factual assertion must include a source reference to a specific uploaded document.
6. Every legal citation must use the EXACT case name and reporter citation from the approved sources. Do not paraphrase or alter citations. For Ohio cases, use the Ohio State Reports citation. For federal cases, use the U.S. Reports or Federal Reporter citation.
7. When citing a legal rule or standard, always include: (a) the specific rule or statute number, (b) the case that established or applied the rule, and (c) the exact citation from the approved sources.
8. This is an EDUCATIONAL tool for law students. Your output is for learning purposes only and must not be filed with any court.
9. If the student asks you to cite something not in the approved sources, explain that you can only work with provided materials and suggest they upload additional documents or add the source to the approved library.
10. For statutes and rules of procedure, always cite the specific section number (e.g., "Fed. R. Civ. P. 56(a)" not just "Rule 56")."""


def build_case_info_section(case_info: dict) -> str:
    """Build the case information section of a prompt."""
    if not case_info or not any(case_info.values()):
        return ""
    parts = ["\n\n--- CASE INFORMATION ---"]
    fields = [
        ("plaintiff", "Plaintiff"),
        ("defendant", "Defendant"),
        ("court", "Court"),
        ("jurisdiction", "Jurisdiction"),
        ("case_number", "Case Number"),
        ("representing_side", "Representing"),
        ("judge", "Judge"),
    ]
    for key, label in fields:
        if case_info.get(key):
            parts.append(f"{label}: {case_info[key]}")
    return "\n".join(parts)


def build_documents_section(documents: list[tuple], max_chars: int = None, priority_type: str = "pleading") -> str:
    """Build the uploaded documents section of a prompt with token budgeting."""
    if not documents:
        return ""
    if max_chars is None:
        max_chars = MAX_DOCUMENT_TOKENS * CHARS_PER_TOKEN
    budgeted_docs = budget_documents(documents, max_chars, priority_type)
    parts = ["\n\n--- UPLOADED DOCUMENTS ---"]
    parts.append("CITATION FORMAT: When citing these documents, use the format: (Last Name, Deposition, at [paragraph]) for depositions, or (Title, at [page]) for exhibits. Do NOT use 'Doc #' references. Use the document TITLE or deponent's LAST NAME.")
    for doc_id, doc_type, title, text in budgeted_docs:
        parts.append(f"\nDocument #{doc_id} ({doc_type}): {title}")
        parts.append(text or "[No text extracted]")
    return "\n".join(parts)


def build_library_section(library_docs: list[tuple]) -> str:
    """Build the library/approved sources section of a prompt."""
    if not library_docs:
        return ""
    parts = ["\n\nLibrary Resources:"]
    budget_per_lib = (MAX_LIBRARY_TOKENS * CHARS_PER_TOKEN) // max(len(library_docs), 1)
    for title, content in library_docs:
        parts.append(f"\n{title}:\n{truncate_text(content, budget_per_lib)}")
    return "\n".join(parts)


# ── DOCX helpers ─────────────────────────────────────────────────────────────

def setup_legal_docx() -> object:
    """Create a new DOCX document with standard legal formatting (TNR 12pt, double-spaced, 1" margins, page numbers)."""
    ensure_docx_imports()
    doc = DocxDocument()

    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        # Page numbers - bottom center
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        fldChar1 = run._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
        run._r.append(fldChar1)
        run2 = fp.add_run()
        instrText = run2._r.makeelement(qn('w:instrText'), {})
        instrText.text = ' PAGE '
        run2._r.append(instrText)
        run3 = fp.add_run()
        fldChar2 = run3._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
        run3._r.append(fldChar2)
        for r in [run, run2, run3]:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    return doc


def add_educational_disclaimer(doc, position: str = "top"):
    """Add educational disclaimer at top or bottom of document."""
    ensure_docx_imports()
    if position == "top":
        add_centered_para(doc, "EDUCATIONAL DOCUMENT — NOT FOR FILING", bold=True, size=11)
        dp = doc.add_paragraph()
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dr = dp.add_run(
            "This document was generated by an AI tool for educational purposes only. "
            "It has not been reviewed by a licensed attorney and should not be filed with any court."
        )
        dr.italic = True
        dr.font.size = Pt(10)
        dr.font.name = "Times New Roman"
        doc.add_paragraph()  # spacer
    else:
        doc.add_paragraph()
        doc.add_paragraph()
        bd = doc.add_paragraph()
        bd.alignment = WD_ALIGN_PARAGRAPH.CENTER
        br = bd.add_run(
            "EDUCATIONAL DOCUMENT — NOT FOR FILING\n"
            "Generated by Sage's Study Group (lawstudygroup.com) for educational purposes only."
        )
        br.italic = True
        br.font.size = Pt(10)
        br.font.name = "Times New Roman"


def build_caption_table(doc, case_info: dict):
    """Build the standard federal court caption block as a table."""
    ensure_docx_imports()

    plaintiff = case_info.get("plaintiff", "[Plaintiff]")
    defendant = case_info.get("defendant", "[Defendant]")
    court = case_info.get("court", "[United States District Court]")
    case_number = case_info.get("case_number", "[Case No.]")

    # Court name
    add_centered_para(doc, "IN THE " + court.upper(), bold=True, size=12)
    doc.add_paragraph()

    # Caption table
    caption_table = doc.add_table(rows=5, cols=3)
    caption_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Remove all borders
    for row in caption_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.makeelement(qn('w:tcBorders'), {})
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = tcBorders.makeelement(qn(f'w:{border_name}'), {
                    qn('w:val'): 'none', qn('w:sz'): '0', qn('w:space'): '0', qn('w:color'): 'auto'
                })
                tcBorders.append(border)
            tcPr.append(tcBorders)

    # Set column widths
    for row in caption_table.rows:
        row.cells[0].width = Inches(3.0)
        row.cells[1].width = Inches(0.5)
        row.cells[2].width = Inches(3.0)

    # Fill in caption
    set_cell_text(caption_table.cell(0, 0), plaintiff + ",", bold=False)
    set_cell_text(caption_table.cell(0, 1), ")", bold=False)
    set_cell_text(caption_table.cell(0, 2), "", bold=False)

    set_cell_text(caption_table.cell(1, 0), "          Plaintiff,", bold=False)
    set_cell_text(caption_table.cell(1, 1), ")", bold=False)
    set_cell_text(caption_table.cell(1, 2), f"Case No. {case_number}", bold=True)

    set_cell_text(caption_table.cell(2, 0), "     v.", bold=False)
    set_cell_text(caption_table.cell(2, 1), ")", bold=False)
    set_cell_text(caption_table.cell(2, 2), "", bold=False)

    set_cell_text(caption_table.cell(3, 0), defendant + ",", bold=False)
    set_cell_text(caption_table.cell(3, 1), ")", bold=False)
    set_cell_text(caption_table.cell(3, 2), "", bold=False)

    set_cell_text(caption_table.cell(4, 0), "          Defendant.", bold=False)
    set_cell_text(caption_table.cell(4, 1), ")", bold=False)
    set_cell_text(caption_table.cell(4, 2), "", bold=False)

    doc.add_paragraph()


def set_cell_text(cell, text: str, bold: bool = False):
    """Set text in a table cell with TNR 12pt formatting."""
    ensure_docx_imports()
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold


def add_centered_para(doc, text: str, bold: bool = False, size: int = 12):
    """Add a centered paragraph."""
    ensure_docx_imports()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return p


def add_signature_block(doc, representing: str, movant: str):
    """Add a right-aligned signature block."""
    ensure_docx_imports()
    lines = [
        "Respectfully submitted,",
        "",
        "_________________________________",
        "[Attorney Name]",
        f"Attorney for {representing.capitalize()}, {movant}",
        "[Address]",
        "[City, State ZIP]",
        "[Phone]",
        "[Bar Number]",
    ]
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_jurat_block(doc, state: str = "___________", county: str = "___________"):
    """Add a jurat/notary block for affidavits."""
    ensure_docx_imports()

    doc.add_paragraph()

    jurat_lines = [
        (f"STATE OF {state.upper()}", False, WD_ALIGN_PARAGRAPH.LEFT),
        ("", False, WD_ALIGN_PARAGRAPH.LEFT),
        (f"COUNTY OF {county.upper()}", False, WD_ALIGN_PARAGRAPH.LEFT),
    ]

    # State/county with "ss:" block
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(0)
    p1.paragraph_format.line_spacing = 1.15
    r = p1.add_run(f"STATE OF {state.upper()}  )")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    spaces = " " * 24
    r = p2.add_run(f"{spaces}) ss:")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.line_spacing = 1.15
    r = p3.add_run(f"COUNTY OF {county.upper()}  )")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

    doc.add_paragraph()

    # Sworn statement
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5)
    r = p.add_run(
        "Subscribed and sworn to before me this _____ day of __________________, ________."
    )
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    # Notary signature
    sig = doc.add_paragraph()
    sig.paragraph_format.space_after = Pt(0)
    r = sig.add_run("_________________________________")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

    name_p = doc.add_paragraph()
    name_p.paragraph_format.space_after = Pt(0)
    r = name_p.add_run("Notary Public")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

    exp_p = doc.add_paragraph()
    r = exp_p.add_run("My commission expires: _______________")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)


def _add_hyperlink(paragraph, url: str, text: str, font_name="Times New Roman", font_size=12):
    """Add a hyperlink to a paragraph. python-docx doesn't support this natively."""
    ensure_docx_imports()
    from docx.oxml import OxmlElement

    # Create the w:hyperlink element
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Blue color + underline for hyperlink style
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '1155CC')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # Font
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size * 2))  # half-points
    rPr.append(sz)

    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    text_elem.set(qn('xml:space'), 'preserve')
    new_run.append(text_elem)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)


# Regex for case reporter citations in generated text
_CASE_CITE_RE = re.compile(r'\b(\d{1,4})\s+([A-Z][A-Za-z0-9.\s\']{1,25}?)\s+(\d{1,5})\b')

SITE_URL = "https://lawstudygroup.com"


def _find_case_citations(text: str, verified_slugs: Optional[set] = None) -> list[tuple[int, int, str]]:
    """Find case citations in text, return list of (start, end, url).

    If verified_slugs is provided, only include citations whose slug is in the set.
    If None, include all valid-looking citations (no DB verification).
    """
    from citation_utils import reporter_cite_to_slug
    results = []
    for m in _CASE_CITE_RE.finditer(text):
        volume, reporter, page = m.group(1), m.group(2).strip(), m.group(3)
        # Skip pin cites like "477 U.S. at 249"
        if re.search(r'\bat$', reporter, re.IGNORECASE):
            continue
        full_cite = f"{volume} {reporter} {page}"
        slug = reporter_cite_to_slug(full_cite)
        # Only link if slug looks like a valid citation
        if re.match(r'^\d+-[a-z].*-\d+$', slug):
            if verified_slugs is not None and slug not in verified_slugs:
                continue
            results.append((m.start(), m.end(), f"{SITE_URL}/cases/{slug}"))
    return results


async def verify_citation_slugs(text: str, db_pool) -> set:
    """Check all detected citation slugs against the database, return set of verified slugs."""
    from citation_utils import reporter_cite_to_slug, parse_citation_slug
    # Collect all unique slugs from the text
    slugs_to_check = {}  # slug -> cite_str
    for m in _CASE_CITE_RE.finditer(text):
        volume, reporter, page = m.group(1), m.group(2).strip(), m.group(3)
        if re.search(r'\bat$', reporter, re.IGNORECASE):
            continue
        full_cite = f"{volume} {reporter} {page}"
        slug = reporter_cite_to_slug(full_cite)
        if re.match(r'^\d+-[a-z].*-\d+$', slug) and slug not in slugs_to_check:
            parsed = parse_citation_slug(slug)
            if parsed:
                v, r, p = parsed
                slugs_to_check[slug] = f"{v} {r} {p}"

    if not slugs_to_check or not db_pool:
        return set()

    verified = set()
    async with db_pool.acquire() as conn:
        for slug, cite_str in slugs_to_check.items():
            cite_upper = cite_str.upper()
            cite_no_dots = re.sub(r'\.(\s?)', ' ', cite_str).strip()
            cite_no_dots = re.sub(r'\s+', ' ', cite_no_dots).upper()
            row = await conn.fetchrow(
                """SELECT id FROM cases
                   WHERE reporter_cite = $1
                      OR reporter_cite LIKE $2
                      OR reporter_cite LIKE $3
                      OR UPPER(reporter_cite) = $4
                      OR UPPER(SPLIT_PART(reporter_cite, ',', 1)) = $4
                      OR UPPER(REPLACE(SPLIT_PART(reporter_cite, ',', 1), '.', ' ')) LIKE $5
                   LIMIT 1""",
                cite_str, f"{cite_str} (%", f"{cite_str},%", cite_upper, f"%{cite_no_dots}%",
            )
            if row:
                verified.add(slug)

    return verified


def add_formatted_run(paragraph, text: str, link_citations: bool = True, verified_slugs: Optional[set] = None):
    """Add a run of text with TNR 12pt to a paragraph, handling inline bold/italic and citation hyperlinks.

    If verified_slugs is provided, only hyperlink citations in that set.
    If None and link_citations is True, hyperlinks all valid-looking citations (no DB check).
    """
    ensure_docx_imports()

    if link_citations:
        citations = _find_case_citations(text, verified_slugs=verified_slugs)
    else:
        citations = []

    if not citations:
        # Original behavior — no citations to link
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                run = paragraph.add_run(part)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        return

    # Split text into citation and non-citation segments
    pos = 0
    for start, end, url in citations:
        # Add text before citation
        if start > pos:
            _add_plain_runs(paragraph, text[pos:start])
        # Add citation as hyperlink
        _add_hyperlink(paragraph, url, text[start:end])
        pos = end
    # Add remaining text
    if pos < len(text):
        _add_plain_runs(paragraph, text[pos:])


# Pattern for case names: "Word v. Word" or "Word v Word" with optional parenthetical
# Matches: "Anderson v. Liberty Lobby, Inc.", "Celotex Corp. v. Catrett", "Mudrich v. Standard Oil Co."
_CASE_NAME_RE = re.compile(
    r'\b([A-Z][A-Za-z\'.]+(?:\s+(?:of|for|the|and|in|ex|re|In|Ex|Re)\s+[A-Z][A-Za-z\'.]+)*'  # plaintiff
    r'(?:\s+[A-Z][A-Za-z\'.]+)*'  # additional plaintiff words
    r'(?:,?\s+(?:Inc|Corp|Co|Ltd|LLC|L\.?L\.?C|P\.?C|S\.?A|N\.?A|et al)\.?)*'  # entity suffixes
    r')\s+v\.?\s+'  # "v." or "v"
    r'([A-Z][A-Za-z\'.]+(?:\s+(?:of|for|the|and|in|ex|re|In|Ex|Re)\s+[A-Z][A-Za-z\'.]+)*'  # defendant
    r'(?:\s+[A-Z][A-Za-z\'.]+)*'  # additional defendant words
    r'(?:,?\s+(?:Inc|Corp|Co|Ltd|LLC|L\.?L\.?C|P\.?C|S\.?A|N\.?A|et al)\.?)*'  # entity suffixes
    r')'
)

# Pattern for "Id." and "Id" (with or without period, standalone)
_ID_RE = re.compile(r'\bId\.(?!\w)|\bId\b(?!\w)')


def _split_on_legal_italics(text: str) -> list[tuple[str, bool]]:
    """Split text into segments that should or shouldn't be italicized.
    Returns list of (text, should_italicize) tuples.
    """
    segments = []
    # Combine case name and Id. patterns
    combined = re.compile(f'({_CASE_NAME_RE.pattern}|{_ID_RE.pattern})')
    pos = 0
    for m in combined.finditer(text):
        if m.start() > pos:
            segments.append((text[pos:m.start()], False))
        segments.append((m.group(0), True))
        pos = m.end()
    if pos < len(text):
        segments.append((text[pos:], False))
    return segments if segments else [(text, False)]


def _add_plain_runs(paragraph, text: str):
    """Add plain text with bold/italic/case-name formatting."""
    ensure_docx_imports()

    # First split on legal italics (case names, Id.)
    for segment, is_legal_italic in _split_on_legal_italics(text):
        # Then handle markdown bold/italic within each segment
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', segment)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
                if is_legal_italic:
                    run.italic = True
            elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                run = paragraph.add_run(part)
                if is_legal_italic:
                    run.italic = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def render_markdown_to_docx(doc, lines: list[str], verified_slugs: Optional[set] = None):
    """
    Render markdown-formatted legal content into DOCX with proper formatting.
    Handles headings (#, ##, ###, ####), numbered lists, and body paragraphs.
    If verified_slugs is provided, only verified citations become hyperlinks.
    """
    ensure_docx_imports()
    for line in lines:
        stripped = line.strip()

        if not stripped:
            continue

        # Section headings (centered, bold, caps)
        if stripped.startswith("# "):
            add_centered_para(doc, stripped[2:].strip().upper(), bold=True, size=12)
            continue

        # Sub-section headings (centered, bold)
        if stripped.startswith("## "):
            add_centered_para(doc, stripped[3:].strip(), bold=True, size=12)
            continue

        # Sub-sub-headings (####) — left-aligned, bold, italic, indented
        if stripped.startswith("#### "):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(stripped[5:].strip())
            run.bold = True
            run.italic = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            continue

        # Argument sub-headings (left-aligned, bold, indented)
        if stripped.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(stripped[4:].strip())
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            continue

        # Bold-only line (often a section label)
        if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            add_centered_para(doc, stripped[2:-2], bold=True, size=12)
            continue

        # Numbered lines (e.g., "1. Fact text here")
        num_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if num_match:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            add_formatted_run(p, f"{num_match.group(1)}. {num_match.group(2)}", verified_slugs=verified_slugs)
            continue

        # Regular body paragraph with first-line indent
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.5)
        add_formatted_run(p, stripped, verified_slugs=verified_slugs)


def save_docx_to_bytes(doc) -> bytes:
    """Save a DOCX document to bytes."""
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ── Citation verification ────────────────────────────────────────────────────

async def verify_citations(text: str, db_pool) -> dict:
    """
    Post-process AI-generated text to verify citations against the database.
    Returns {verified: [...], unverified: [...], text: modified_text}
    """
    try:
        from eyecite import get_citations
        from eyecite.models import CaseCitation
    except ImportError:
        return {"verified": [], "unverified": [], "text": text}

    citations = get_citations(text)
    verified = []
    unverified = []

    async with db_pool.acquire() as conn:
        for cite in citations:
            if not isinstance(cite, CaseCitation):
                continue

            cite_str = str(cite)
            row = await conn.fetchrow(
                "SELECT id, title FROM cases WHERE reporter_cite ILIKE $1 LIMIT 1",
                f"%{cite_str}%"
            )
            if row:
                verified.append({"citation": cite_str, "case_id": row["id"], "title": row["title"]})
            else:
                unverified.append({"citation": cite_str})

    return {
        "verified": verified,
        "unverified": unverified,
        "text": text,
    }
