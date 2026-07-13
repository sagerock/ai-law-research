#!/usr/bin/env python3
"""Build the in-site reader content for a casebook from its source .docx.

Walks the document in order (paragraphs + tables), maps each Word style to a
semantic block, groups consecutive rule/note/problem/quote/diagram runs into one
container, and assigns every paragraph a stable anchor `para-<n>` where <n> is the
docx paragraph index. That index matches the `para` field in the casebook's Qdrant
collection, so AI Q&A citations can deep-link to the exact paragraph in the reader.

Usage:
    python scripts/build_casebook_reader.py --docx PATH [--casebook-id 2467]
    python scripts/build_casebook_reader.py --docx PATH --dry-run --html out.html

--dry-run renders every chapter to a single HTML file and prints stats; it does not
touch the database. Without --dry-run it loads rows into casebook_content (replacing
any existing rows for that casebook).
"""
import argparse
import html as htmllib
import os
import re
import sys

import docx
from docx.text.paragraph import Paragraph
from docx.table import Table

# Word style name -> semantic block type. Chapter/Book styles are handled
# separately (they set structure); everything else maps here, default "text".
STYLE_TO_TYPE = {
    "Case": "case", "Judge": "judge", "Rule": "rule", "Notes": "note",
    "Problem": "problem", "Quote Block": "quote", "block quote": "quote",
    "Hearsay Diagram": "diagram", "Divider": "divider", "List Paragraph": "list",
}
# In the source, Rule/Notes/Problem styles mark a *label* line; the body that
# follows is plain Text. So a label opens a boxed group that absorbs the
# following body blocks until the next structural break (heading/case/etc.).
LABEL_TYPES = {"rule", "note", "problem"}
ABSORB_TYPES = {"text", "quote", "list", "diagram", "table", "image"}


def render_chain(para) -> str:
    """A 'Hearsay Diagram' paragraph holds node labels (e.g. Hillmon / Letter /
    Sister) as separate runs. Render them as a left-to-right box-and-arrow chain."""
    nodes = [r.text.strip() for r in para.runs if r.text and r.text.strip()]
    if not nodes:
        return ""
    parts = []
    for k, n in enumerate(nodes):
        if k:
            parts.append('<span class="darrow">→</span>')
        parts.append(f'<span class="dnode">{htmllib.escape(n)}</span>')
    return "".join(parts)


def chapter_slug(title: str) -> str:
    m = re.match(r"\s*CHAPTER\s+(\d+)", title, re.I)
    if m:
        return f"ch{m.group(1)}"
    return "front-" + re.sub(r"[^a-z0-9]+", "-", title.lower()).strip("-")


def render_inline(para) -> str:
    """Inner HTML for a paragraph, preserving italic/bold runs."""
    if not para.runs:
        return htmllib.escape(para.text)
    out = []
    for r in para.runs:
        t = htmllib.escape(r.text or "")
        if not t:
            continue
        if r.italic:
            t = f"<em>{t}</em>"
        if r.bold:
            t = f"<strong>{t}</strong>"
        out.append(t)
    return "".join(out) or htmllib.escape(para.text)


def render_table(table) -> str:
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            txt = "<br>".join(htmllib.escape(p.text) for p in cell.paragraphs if p.text.strip())
            cells.append(f"<td>{txt}</td>")
        rows.append("<tr>" + "".join(cells) + "</tr>")
    return "<table class='cb-table'>" + "".join(rows) + "</table>"


def parse(docx_path: str, casebook_id: int = 2467):
    """Return (blocks, chapters, images). Each block is a dict ready for
    casebook_content; images is the set of media filenames to extract."""
    import zipfile
    z = zipfile.ZipFile(docx_path)
    rels = z.read("word/_rels/document.xml.rels").decode()
    rid_to_media = {m[0]: m[1].split("/")[-1]
                    for m in re.findall(r'Id="([^"]+)"[^>]*Target="(media/[^"]+)"', rels)}
    img_base = f"/casebook/{casebook_id}"
    images = []  # referenced media filenames, in order

    doc = docx.Document(docx_path)
    blocks = []
    chapters = []  # ordered [(slug, title)]
    cur_slug, cur_title, cur_section = None, None, None
    skipping = False            # inside TABLE OF CONTENTS
    group_counter = 0
    active_gid = None           # open boxed group (started by a label)
    active_label = None
    sort_order = 0
    p_idx = -1                  # index into doc.paragraphs (== Qdrant 'para')

    def emit(block_type, html, para_ordinal=None):
        nonlocal sort_order, group_counter, active_gid, active_label
        if block_type in LABEL_TYPES:
            if active_gid is not None and block_type == active_label:
                gid = active_gid                    # merge consecutive same labels
            else:
                group_counter += 1
                active_gid, active_label, gid = group_counter, block_type, group_counter
        elif block_type in ABSORB_TYPES:
            gid = active_gid                         # join the open box, if any
        else:
            active_gid, active_label, gid = None, None, None   # structural break
        blocks.append({
            "sort_order": sort_order,
            "chapter_slug": cur_slug, "chapter_title": cur_title, "section": cur_section,
            "block_type": block_type, "group_id": gid, "html": html,
            "para_ordinal": para_ordinal,
            "anchor": f"para-{para_ordinal}" if para_ordinal is not None else None,
        })
        sort_order += 1

    for child in doc.element.body.iterchildren():
        tag = child.tag
        if tag.endswith("}p"):
            p_idx += 1
            para = Paragraph(child, doc)
            style = para.style.name if para.style else "Normal"

            if style.startswith("toc"):
                continue
            if style == "Chapter":
                title = para.text.strip()
                if not title:
                    continue
                if title.upper().startswith("TABLE OF CONTENTS"):
                    skipping = True
                    continue
                skipping = False
                cur_slug, cur_title, cur_section = chapter_slug(title), title, None
                prev_group_type = None
                chapters.append((cur_slug, cur_title))
                emit("chapter-title", htmllib.escape(title), p_idx)
                continue
            if skipping or cur_slug is None:
                continue
            if style == "Book 1":
                cur_section = para.text.strip()
                emit("section", htmllib.escape(cur_section), p_idx)
                continue
            if style == "Book 2":
                emit("subsection", render_inline(para), p_idx)
                continue

            # Embedded figures (testimonial triangle, timelines, exhibits) live in
            # otherwise-empty paragraphs as drawing/picture refs. Emit them as images.
            for rid in re.findall(r'(?:r:embed|r:id)="([^"]+)"', para._p.xml):
                fn = rid_to_media.get(rid)
                if not fn:
                    continue  # non-image rel (e.g. a hyperlink)
                if fn not in images:
                    images.append(fn)
                emit("image",
                     f'<img class="cb-figure-img" src="{img_base}/{fn}" '
                     f'alt="Figure from {cur_title.strip()}" loading="lazy">', p_idx)

            if style == "Hearsay Diagram":
                html = render_chain(para)
                if html:
                    emit("diagram", html, p_idx)
                continue

            block_type = STYLE_TO_TYPE.get(style, "text")
            if block_type == "divider":
                emit("divider", "", p_idx)
                continue
            html = render_inline(para)
            if not html.strip():
                continue  # drop empty spacer paragraphs
            emit(block_type, html, p_idx)

        elif tag.endswith("}tbl"):
            if skipping or cur_slug is None:
                continue
            emit("table", render_table(Table(child, doc)))

    return blocks, chapters, images


def extract_images(docx_path, images, out_dir):
    """Copy the referenced media files out of the .docx into out_dir."""
    import zipfile
    if not images:
        return
    os.makedirs(out_dir, exist_ok=True)
    z = zipfile.ZipFile(docx_path)
    for fn in images:
        with open(os.path.join(out_dir, fn), "wb") as f:
            f.write(z.read(f"word/media/{fn}"))


# --------------------------------------------------------------------------- #
# Dry-run HTML rendering (visual QA, no DB)
# --------------------------------------------------------------------------- #
TAG = {"chapter-title": "h1", "section": "h2", "subsection": "h3", "case": "h3",
       "judge": "p", "rule": "div", "note": "div", "problem": "div",
       "quote": "blockquote", "diagram": "pre", "list": "li", "text": "p"}
CSS = """
body{background:#f7f5f1;margin:0;font-family:Georgia,serif;color:#1c1917;line-height:1.7}
.wrap{max-width:720px;margin:0 auto;padding:40px 24px 120px}
h1{font-family:system-ui,sans-serif;font-size:13px;letter-spacing:.12em;text-transform:uppercase;color:#6b7f5e;border-bottom:2px solid #cfd8c4;padding-bottom:12px;margin:56px 0 24px}
h2{font-family:system-ui,sans-serif;font-size:26px;margin:44px 0 8px}
h3{font-family:system-ui,sans-serif;font-size:18px;margin:30px 0 6px}
.case{padding-top:10px;border-top:1px solid #e6e0d6}
.judge{font-style:italic;color:#57534e;margin:.2em 0 1em}
p{margin:0 0 1.05em}
.rule{background:#eef1e9;border-left:3px solid #8aa173;padding:6px 16px;border-radius:4px;margin:1em 0}
.note{background:#faf8f3;border-left:2px solid #d8cdb8;padding:6px 16px;margin:1em 0;font-size:.96em}
.problem{background:#fbf3e4;border:1px solid #ecd9b4;padding:8px 16px;border-radius:6px;margin:1em 0}
.rule p,.note p,.problem p{margin:.55em 0}
.rule .label,.note .label,.problem .label{font-family:system-ui,sans-serif;font-size:11px;letter-spacing:.08em;text-transform:uppercase;font-weight:600;color:#6b7f5e;margin:.2em 0 .5em}
.note .label{color:#a08a5b}.problem .label{color:#b07d3a}
.rule blockquote,.note blockquote,.problem blockquote{margin:.5em 0}
blockquote{margin:1em 0;padding:2px 0 2px 20px;border-left:3px solid #d6cdbd;color:#37322c;font-size:.97em}
.cb-diagram{display:flex;flex-wrap:wrap;align-items:center;gap:8px;margin:1.4em 0;font-family:system-ui,sans-serif}
.dnode{border:1.5px solid #8aa173;background:#eef1e9;border-radius:6px;padding:6px 12px;font-size:.9em;font-weight:600;color:#2b2520}
.darrow{color:#8aa173;font-size:1.2em}
.cb-figure{margin:1.6em 0;text-align:center}
.cb-figure-img{max-width:100%;height:auto;border:1px solid #e6e0d6;border-radius:6px}
.cb-table{border-collapse:collapse;margin:1em 0;font-size:.9em}
.cb-table td{border:1px solid #cfc8ba;padding:6px 10px;vertical-align:top}
.license{margin-top:64px;padding-top:18px;border-top:1px solid #ddd;font-family:sans-serif;font-size:12px;color:#78716c}
.license a{color:#6b7f5e}
"""


def render_html(blocks) -> str:
    parts, i = [], 0
    while i < len(blocks):
        b = blocks[i]
        gid, btype = b["group_id"], b["block_type"]
        anchor = f' id="{b["anchor"]}"' if b["anchor"] else ""
        if gid is not None:
            run = [b]
            while i + 1 < len(blocks) and blocks[i + 1]["group_id"] == gid:
                i += 1
                run.append(blocks[i])
            inner = []
            for x in run:
                a = f' id="{x["anchor"]}"' if x["anchor"] else ""
                t = x["block_type"]
                if t in ("rule", "note", "problem"):
                    inner.append(f'<p{a} class="label">{x["html"]}</p>')
                elif t == "quote":
                    inner.append(f"<blockquote{a}>{x['html']}</blockquote>")
                elif t == "table":
                    inner.append(x["html"])
                elif t == "diagram":
                    inner.append(f"<div{a} class='cb-diagram'>{x['html']}</div>")
                elif t == "image":
                    inner.append(f"<figure{a} class='cb-figure'>{x['html']}</figure>")
                else:
                    inner.append(f"<p{a}>{x['html']}</p>")
            parts.append(f'<div class="{run[0]["block_type"]}">{"".join(inner)}</div>')
        elif btype == "divider":
            parts.append(f"<hr{anchor}>")
        elif btype == "table":
            parts.append(b["html"])
        elif btype == "diagram":
            parts.append(f"<div{anchor} class='cb-diagram'>{b['html']}</div>")
        elif btype == "image":
            parts.append(f"<figure{anchor} class='cb-figure'>{b['html']}</figure>")
        else:
            tag = TAG.get(btype, "p")
            cls = f' class="{btype}"' if btype in ("case", "judge") else ""
            parts.append(f"<{tag}{anchor}{cls}>{b['html']}</{tag}>")
        i += 1
    return (f"<!doctype html><html><head><meta charset='utf-8'>"
            f"<meta name='viewport' content='width=device-width,initial-scale=1'>"
            f"<style>{CSS}</style></head><body><div class='wrap'>"
            + "\n".join(parts)
            + "<div class='license'>Evidence by Edward K. Cheng — reformatted for the web by "
              "Tortwell and licensed under "
              "<a href='https://creativecommons.org/licenses/by-nc-sa/4.0/'>CC BY-NC-SA 4.0</a>.</div>"
            + "</div></body></html>")


def load_db(blocks, casebook_id):
    import psycopg2
    from psycopg2.extras import execute_values
    dsn = (os.getenv("DATABASE_PUBLIC_URL")
           or os.getenv("DATABASE_URL", "").replace(
               "postgres.railway.internal:5432", "switchyard.proxy.rlwy.net:22438"))
    if not dsn:
        sys.exit("No DATABASE_PUBLIC_URL / DATABASE_URL set.")
    conn = psycopg2.connect(dsn)
    cur = conn.cursor()
    cur.execute("DELETE FROM casebook_content WHERE casebook_id = %s", (casebook_id,))
    execute_values(cur, """
        INSERT INTO casebook_content
          (casebook_id, sort_order, chapter_slug, chapter_title, section,
           block_type, group_id, html, para_ordinal, anchor)
        VALUES %s
    """, [(casebook_id, b["sort_order"], b["chapter_slug"], b["chapter_title"], b["section"],
           b["block_type"], b["group_id"], b["html"], b["para_ordinal"], b["anchor"])
          for b in blocks])
    conn.commit()
    n = cur.rowcount
    conn.close()
    return len(blocks)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--docx", required=True)
    ap.add_argument("--casebook-id", type=int, default=2467)
    ap.add_argument("--dry-run", action="store_true")
    ap.add_argument("--html", default="reader_preview.html")
    ap.add_argument("--image-dir", default=None,
                    help="where to write extracted figures (default frontend/public/casebook/<id>)")
    args = ap.parse_args()

    blocks, chapters, images = parse(args.docx, args.casebook_id)
    image_dir = args.image_dir or f"frontend/public/casebook/{args.casebook_id}"

    from collections import Counter
    types = Counter(b["block_type"] for b in blocks)
    print(f"parsed {len(blocks)} blocks across {len(chapters)} chapters")
    print("chapters:")
    for slug, title in chapters:
        n = sum(1 for b in blocks if b["chapter_slug"] == slug)
        print(f"  {slug:14} {n:5d} blocks  {title}")
    print("block types:", dict(types))
    print(f"figures: {len(images)} -> {image_dir}  ({', '.join(images)})")

    if args.dry_run:
        # extract into the preview dir so <img> src paths resolve when served locally
        extract_images(args.docx, images, os.path.join(os.path.dirname(args.html) or ".",
                                                        "casebook", str(args.casebook_id)))
        with open(args.html, "w") as f:
            f.write(render_html(blocks))
        print(f"\nwrote preview -> {args.html}")
    else:
        extract_images(args.docx, images, image_dir)
        n = load_db(blocks, args.casebook_id)
        print(f"\nloaded {n} rows into casebook_content for casebook {args.casebook_id}")


if __name__ == "__main__":
    main()
